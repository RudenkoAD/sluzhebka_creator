from docx import Document
from datetime import datetime
def gen_sluzhebka(group, genitive_name, event_name, date, start_time, end_time, num_people, name):
  if group == "" or genitive_name == "" or event_name == "" or date == "" or start_time == "" or end_time == "" or num_people == "" or name == "":
    raise Exception("Empty field")
  doc = Document('Sluzhebka_bron_12.docx')
  doc.paragraphs[2].text = f"от студента(ки) группы {group}"
  doc.paragraphs[3].text = f"{genitive_name}"
  #transform date from YYYY-MM-DD to DD.MM.YY
  date = date.split("-")
  date = date[2] + "." + date[1] + "." + date[0][2:] 
  doc.paragraphs[14].text = f"Прошу разрешить провести мероприятие “{event_name}” в клубе общежития №6 {date} с {start_time} до {end_time}.  Предполагаемое количество людей – {num_people}."
  doc.paragraphs[20].text = f"{name}."
  format = "%d_%m_%Y_%H_%M_%S  "
  time1 = datetime.now().strftime(format)
  fpath = f"Uploads/Sluzhebka_{time1}.docx"
  doc.save(fpath)
  return fpath

if __name__ == "__main__":
  gen_sluzhebka("Б02-108", "Иванова Ивана Ивановича", "День рождения", "12.12.2021", "12:00", "13:00", "100", "Иванов Иван Иванович")